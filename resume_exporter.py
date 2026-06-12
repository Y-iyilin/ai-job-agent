from io import BytesIO
from pathlib import Path
import re
from zipfile import ZIP_DEFLATED, ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraph
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph as PdfParagraph
from reportlab.platypus import SimpleDocTemplate, Spacer
from pypdf import PdfReader


RESUME_TEMPLATES = {
    "简洁商务": {
        "description": "适合大多数应届生和职能类岗位，排版克制、通用、稳妥。",
        "font_size": 10.5,
        "title_size": 16,
        "heading_size": 13,
        "heading_color": "2F5597",
        "margin": 0.7,
        "pdf_heading_color": colors.HexColor("#2F5597"),
    },
    "技术清爽": {
        "description": "适合技术支持、测试、数据分析、开发实习等偏技术岗位。",
        "font_size": 10,
        "title_size": 15,
        "heading_size": 12.5,
        "heading_color": "1F4E79",
        "margin": 0.65,
        "pdf_heading_color": colors.HexColor("#1F4E79"),
    },
    "实施售前": {
        "description": "适合软件实施、售前技术支持、解决方案、客户成功等偏交付岗位。",
        "font_size": 10.5,
        "title_size": 16,
        "heading_size": 13,
        "heading_color": "548235",
        "margin": 0.7,
        "pdf_heading_color": colors.HexColor("#548235"),
    },
    "应届生一页": {
        "description": "适合内容不多、希望压缩成一页的应届生简历。",
        "font_size": 9.5,
        "title_size": 14,
        "heading_size": 11.5,
        "heading_color": "404040",
        "margin": 0.55,
        "pdf_heading_color": colors.HexColor("#404040"),
    },
}

WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
ET.register_namespace("w", WORD_NS["w"])


def get_template_names() -> list[str]:
    return list(RESUME_TEMPLATES.keys())


def get_template_description(template_name: str) -> str:
    return str(_template(template_name).get("description", "通用简历导出模板。"))


def build_template_preview_html(template_name: str) -> str:
    tpl = _template(template_name)
    color = str(tpl["heading_color"])
    font_size = float(tpl["font_size"])
    return f"""
    <div style="border:1px solid #e5e7eb;border-radius:8px;padding:14px 16px;background:#ffffff;max-width:520px;margin:6px 0 12px 0;">
      <div style="font-size:{font_size + 6}px;font-weight:700;color:#{color};margin-bottom:6px;">张三</div>
      <div style="font-size:12px;color:#4b5563;margin-bottom:10px;">软件工程应届生｜杭州｜138****0000｜email@example.com</div>
      <div style="height:1px;background:#{color};opacity:.55;margin:8px 0 10px 0;"></div>
      <div style="font-size:{font_size + 2}px;font-weight:700;color:#{color};margin-top:8px;">求职方向</div>
      <div style="font-size:{font_size}px;color:#111827;line-height:1.7;">软件实施 / 技术支持 / 产品助理</div>
      <div style="font-size:{font_size + 2}px;font-weight:700;color:#{color};margin-top:8px;">项目经历</div>
      <ul style="font-size:{font_size}px;color:#111827;line-height:1.7;margin-top:4px;">
        <li>学生管理系统：负责数据库设计、功能测试和项目文档整理。</li>
        <li>使用 AI 工具辅助整理需求说明和培训材料。</li>
      </ul>
    </div>
    """


def _template(template_name: str) -> dict:
    return RESUME_TEMPLATES.get(template_name, RESUME_TEMPLATES["简洁商务"])


def _hex_to_rgb_color(value: str) -> RGBColor:
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _clean_markdown_line(line: str) -> str:
    line = line.strip()
    line = re.sub(r"^\s{0,3}#{1,6}\s*", "", line)
    line = re.sub(r"^\s*[-*]\s*", "• ", line)
    line = line.replace("`", "")
    line = line.replace("**", "")
    return line.strip()


def _iter_lines(markdown_text: str) -> list[tuple[str, str]]:
    lines = []
    for raw_line in markdown_text.splitlines():
        if not raw_line.strip():
            continue
        stripped = raw_line.strip()
        if stripped.startswith("## "):
            lines.append(("heading", _clean_markdown_line(stripped)))
        elif stripped.startswith("# "):
            lines.append(("title", _clean_markdown_line(stripped)))
        elif stripped.startswith(("-", "*")):
            lines.append(("bullet", _clean_markdown_line(stripped)))
        else:
            lines.append(("body", _clean_markdown_line(stripped)))
    return lines


def _insert_paragraph_after(paragraph: DocxParagraph, text: str = "", style: str | None = None) -> DocxParagraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = DocxParagraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _replace_paragraph_text(paragraph: DocxParagraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _template_paragraphs(document: Document) -> list[DocxParagraph]:
    paragraphs: list[DocxParagraph] = []
    paragraphs.extend(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _resume_lines_for_template(markdown_text: str) -> list[str]:
    lines = []
    for kind, text in _iter_lines(markdown_text):
        if not text:
            continue
        if kind == "bullet":
            lines.append(text)
        else:
            lines.append(text)
    return lines


def _replace_docx_ooxml_text(template_bytes: bytes, resume_lines: list[str]) -> bytes | None:
    if not resume_lines:
        return None

    source = BytesIO(template_bytes)
    output = BytesIO()
    replaced_count = 0
    line_index = 0
    last_text_node: ET.Element | None = None
    xml_cache: dict[str, bytes] = {}

    with ZipFile(source, "r") as zin:
        xml_names = [
            name
            for name in zin.namelist()
            if name == "word/document.xml"
            or name.startswith("word/header")
            or name.startswith("word/footer")
        ]
        for name in xml_names:
            raw_xml = zin.read(name)
            try:
                root = ET.fromstring(raw_xml)
            except ET.ParseError:
                continue
            changed = False
            for paragraph in root.findall(".//w:p", WORD_NS):
                text_nodes = paragraph.findall(".//w:t", WORD_NS)
                if not text_nodes:
                    continue
                original_text = "".join(node.text or "" for node in text_nodes).strip()
                if not original_text:
                    continue
                replacement = resume_lines[line_index] if line_index < len(resume_lines) else ""
                text_nodes[0].text = replacement
                for node in text_nodes[1:]:
                    node.text = ""
                last_text_node = text_nodes[0]
                line_index += 1
                replaced_count += 1
                changed = True
            if changed and last_text_node is not None and line_index < len(resume_lines):
                remaining = "\n".join(resume_lines[line_index:])
                last_text_node.text = f"{last_text_node.text or ''}\n{remaining}".strip()
                line_index = len(resume_lines)
            if changed:
                xml_cache[name] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

        if replaced_count == 0:
            return None

        with ZipFile(output, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = xml_cache.get(item.filename)
                if data is None:
                    data = zin.read(item.filename)
                zout.writestr(item, data)

    return output.getvalue()


def extract_resume_document(markdown_text: str) -> str:
    markers = [
        "## 九、可直接复制的简历版本",
        "## 9、可直接复制的简历版本",
        "可直接复制的简历版本",
    ]
    for marker in markers:
        index = markdown_text.find(marker)
        if index >= 0:
            content = markdown_text[index + len(marker):].strip()
            return content or markdown_text
    return markdown_text


def build_docx(markdown_text: str, title: str = "优化版简历", template_name: str = "简洁商务") -> bytes:
    tpl = _template(template_name)
    document = Document()
    section = document.sections[0]
    margin = float(tpl["margin"])
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)

    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(float(tpl["font_size"]))
    styles["Heading 1"].font.name = "Microsoft YaHei"
    styles["Heading 1"].font.size = Pt(float(tpl["title_size"]))
    styles["Heading 1"].font.color.rgb = _hex_to_rgb_color(str(tpl["heading_color"]))
    styles["Heading 2"].font.name = "Microsoft YaHei"
    styles["Heading 2"].font.size = Pt(float(tpl["heading_size"]))
    styles["Heading 2"].font.color.rgb = _hex_to_rgb_color(str(tpl["heading_color"]))

    document.add_heading(title, level=1)
    for kind, text in _iter_lines(markdown_text):
        if not text:
            continue
        if kind == "title":
            document.add_heading(text, level=1)
        elif kind == "heading":
            document.add_heading(text, level=2)
        elif kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(text.replace("• ", "", 1))
        else:
            document.add_paragraph(text)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_docx_from_template(
    template_bytes: bytes,
    markdown_text: str,
    title: str = "AI 简历优化稿",
    template_name: str = "简洁商务",
) -> bytes:
    tpl = _template(template_name)
    document = Document(BytesIO(template_bytes))
    placeholder_found = False
    for paragraph in document.paragraphs:
        if "{{AI_RESUME}}" in paragraph.text or "{{简历正文}}" in paragraph.text:
            placeholder_found = True
            paragraph.text = title
            anchor = paragraph
            for kind, text in _iter_lines(markdown_text):
                if not text:
                    continue
                if kind in {"title", "heading"}:
                    anchor = _insert_paragraph_after(anchor, text, style="Heading 2")
                    for run in anchor.runs:
                        run.font.color.rgb = _hex_to_rgb_color(str(tpl["heading_color"]))
                elif kind == "bullet":
                    anchor = _insert_paragraph_after(anchor, text.replace("• ", "", 1), style="List Bullet")
                else:
                    anchor = _insert_paragraph_after(anchor, text)
            break

    if placeholder_found:
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    resume_lines = _resume_lines_for_template(markdown_text)
    replaced_docx = _replace_docx_ooxml_text(template_bytes, resume_lines)
    if replaced_docx:
        return replaced_docx

    editable_paragraphs = [
        paragraph
        for paragraph in _template_paragraphs(document)
        if paragraph.text.strip()
    ]
    if editable_paragraphs:
        for index, paragraph in enumerate(editable_paragraphs):
            replacement = resume_lines[index] if index < len(resume_lines) else ""
            _replace_paragraph_text(paragraph, replacement)
        anchor = editable_paragraphs[-1]
        for extra_line in resume_lines[len(editable_paragraphs):]:
            anchor = _insert_paragraph_after(anchor, extra_line, style=anchor.style)
    else:
        document.add_heading(title, level=1)
        for line in resume_lines:
            document.add_paragraph(line)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def extract_docx_template_text(template_bytes: bytes, max_chars: int = 5000) -> str:
    document = Document(BytesIO(template_bytes))
    lines = []
    for paragraph in _template_paragraphs(document):
        text = paragraph.text.strip()
        if text:
            lines.append(text)
        if sum(len(item) for item in lines) >= max_chars:
            break
    return "\n".join(lines)[:max_chars]


def extract_pdf_template_text(template_bytes: bytes, max_chars: int = 5000) -> str:
    reader = PdfReader(BytesIO(template_bytes))
    lines = []
    for page in reader.pages[:5]:
        text = page.extract_text() or ""
        for line in text.splitlines():
            cleaned = line.strip()
            if cleaned:
                lines.append(cleaned)
            if sum(len(item) for item in lines) >= max_chars:
                break
        if sum(len(item) for item in lines) >= max_chars:
            break
    return "\n".join(lines)[:max_chars]


def preview_docx_template(template_bytes: bytes) -> dict[str, str | int | list[str]]:
    document = Document(BytesIO(template_bytes))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_count = len(document.tables)
    image_count = len(document.inline_shapes)
    section_count = len(document.sections)
    return {
        "paragraph_count": len(paragraphs),
        "table_count": table_count,
        "image_count": image_count,
        "section_count": section_count,
        "sample_text": paragraphs[:8],
    }


def preview_pdf_template(template_bytes: bytes) -> dict[str, int | list[str]]:
    reader = PdfReader(BytesIO(template_bytes))
    sample_text = []
    for page in reader.pages[:2]:
        text = page.extract_text() or ""
        for line in text.splitlines():
            cleaned = line.strip()
            if cleaned:
                sample_text.append(cleaned)
            if len(sample_text) >= 8:
                break
        if len(sample_text) >= 8:
            break
    return {
        "page_count": len(reader.pages),
        "sample_text": sample_text[:8],
    }


def _register_chinese_font() -> str:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
    ]
    for path in candidates:
        if path.exists():
            try:
                pdfmetrics.registerFont(TTFont("ResumeChinese", str(path)))
                return "ResumeChinese"
            except Exception:
                continue
    return "Helvetica"


def build_pdf(markdown_text: str, title: str = "优化版简历", template_name: str = "简洁商务") -> bytes:
    tpl = _template(template_name)
    font_name = _register_chinese_font()
    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )
    base_styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=base_styles["Title"],
        fontName=font_name,
        fontSize=float(tpl["title_size"]) + 2,
        leading=24,
        spaceAfter=10,
        textColor=tpl["pdf_heading_color"],
    )
    heading_style = ParagraphStyle(
        "ResumeHeading",
        parent=base_styles["Heading2"],
        fontName=font_name,
        fontSize=float(tpl["heading_size"]),
        leading=18,
        spaceBefore=8,
        spaceAfter=4,
        textColor=tpl["pdf_heading_color"],
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=base_styles["BodyText"],
        fontName=font_name,
        fontSize=float(tpl["font_size"]),
        leading=16,
        spaceAfter=4,
    )

    story = [PdfParagraph(title, title_style)]
    for kind, text in _iter_lines(markdown_text):
        if not text:
            continue
        escaped = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if kind in {"title", "heading"}:
            story.append(PdfParagraph(escaped, heading_style))
        else:
            story.append(PdfParagraph(escaped, body_style))
            story.append(Spacer(1, 2))

    doc.build(story)
    return output.getvalue()
