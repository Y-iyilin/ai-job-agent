from io import BytesIO

from docx import Document
from pypdf import PdfReader


class ResumeParseError(RuntimeError):
    pass


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ResumeParseError("TXT 文件编码无法识别，请尝试另存为 UTF-8 后再上传。")


def _parse_docx(data: bytes) -> str:
    document = Document(BytesIO(data))
    parts = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def parse_resume_file(filename: str, data: bytes) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if suffix == "txt":
        text = _decode_text(data)
    elif suffix == "docx":
        text = _parse_docx(data)
    elif suffix == "pdf":
        text = _parse_pdf(data)
    else:
        raise ResumeParseError("暂时只支持 .txt、.docx、.pdf 格式。")

    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not text:
        raise ResumeParseError("没有识别到可用文字。如果是扫描版 PDF，需要后续接入 OCR。")
    return text[:12000]
